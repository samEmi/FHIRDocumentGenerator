import requests
from docx import Document
from math import floor
import pandas as pd
import plotly.express as px
class Model:
    def __init__(self, filePath, endpoint='http://127.0.0.1:5002/api/patients?'):
        self.data = None
        self.endpoint = endpoint
        self.filePath = filePath
        self.document = Document()

    def getAgeToObservation(self):
        try:
            self.data = requests.get(self.endpoint + "graph=1")
            self.data = self.data.json()
        except:
            raise Exception("Cannot Plot Graph")

        self.plotGraph()

    def plotGraph(self):
        table = []
        for element in self.data:
            ageFreq = self.data[element]
            ageFreq = {int(k): v for k, v in ageFreq.items()}
            for age in sorted(ageFreq.keys()):
                row = []
                frequency = ageFreq[age]
                row.append(age)
                row.append(frequency)
                row.append(element)
                table.append(row)
        self.drawMethodComparisons(table)

    def drawMethodComparisons(self, table):
        df = pd.DataFrame(table)
        df.columns = ["Age", "Frequency", "Category"]
        fig = px.line(df, x="Age", y="Frequency", color="Category", hover_data=["Category"])
        fig.show()


    def makeDocument(self, id):
        self.__sanitise(id)
        try:
            self.data = requests.get(self.endpoint +"id=" + id)
            self.data = self.data.json()
        except Exception:
            raise Exception("Unable retrieve patient data. Please try another ID")
        self.__createDocument()

    def __createDocument(self):
        self.document = Document()
        self.document.add_heading('Patient Observations', 1)
        self.document.add_paragraph('This Document provides all vital sign observations to date for '
                                + self.data["full_name"] + " starting with the latest observation")
        p2 = self.document.add_paragraph("Age:   " + str(floor(self.data["age"])) + " years old            ")
        p2.add_run("Gender:   " + self.data["gender"])
        self.__makeTable()
        self.document.save(self.filePath + ".docx")



    def __makeTable(self):
        table = self.document.add_table(rows=1, cols=5)
        colCells = table.rows[0].cells
        colCells[0].text = 'Number'
        colCells[1].text = 'Type'
        colCells[2].text = 'Status'
        colCells[3].text = 'Date and Time'
        colCells[4].text = 'Components'
        self.__addRows(table)

    def __addRows(self, table):
        self.data["observations"].reverse()
        for i, observation in enumerate(self.data["observations"]):
            if observation["type"] == "vital-signs":
                rowCells = table.add_row().cells
                rowCells[0].text = str(i)
                rowCells[1].text = observation["type"]
                rowCells[2].text = observation["status"]
                rowCells[3].text = observation["date"]
                rowCells[4].text = observation["components"]

    def __sanitise(self, id):
        if not id:
            raise ValueError("ID should not be empty")

